from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from django.conf import settings


def generer_docx(context):
    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = Pt(16)
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    # Contenu

    # # N'affiche pas le logo'
    # logo_path = os.path.join(settings.STATICFILES_DIRS[0], 'images', 'axa_logo.png')
    # if os.path.exists(logo_path):
    #     doc.add_picture(logo_path, width=Inches(1.5))
    #     # last_paragraph = doc.paragraphs[-1]
    #     # last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    run = p.add_run("TARIFICATION INDICATIVE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x66)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Produit chantier")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x66)

    doc.add_paragraph("Tarification indicative (*) sur la base d’un risque conforme aux paramètres suivants :")
    doc.add_paragraph(f"Type d’ouvrage : {context['type_ouvrage']}")
    doc.add_paragraph(f"Types de travaux réalisés : {context['types_travaux']}")
    doc.add_paragraph(f"Coût du chantier : {context['cout_chantier']} €")
    doc.add_paragraph(f"Présence d’existant : {context['existant']}")
    doc.add_paragraph(f"Garantie choisie : {context['garantie']}")
    doc.add_paragraph(f"Description de l’ouvrage : {context['description']}")
    doc.add_paragraph(f"Adresse du chantier : {context['adresse']}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Garanties Tous Risques chantier")
    run.underline = True

    doc.add_paragraph("MONTANTS DE GARANTIES (exprimés en €)").runs[0].bold = True

    lignes1 = [
        ["Dommages matériels à l'ouvrage\n(Ce montant est épuisable pendant la durée des travaux)", context['dommages_ouvrage']],
        ["Responsabilité civile (tous dommages confondus)\n(Ces montants sont épuisables pendant la durée des travaux)", context['responsabilite_civile']],
        ["Maintenance-visite\n(Ce montant est compris dans le montant de la garantie des dommages matériels à l'ouvrage)", context['maintenance']],
        ["Mesure conservatoire", context['mesure_conservatoire']],
    ]

    table1 = doc.add_table(rows=0, cols=2)
    table1.style = 'Table Grid'
    widths = [0.7, 0.3]

    for libelle, valeur in lignes1:
        row_cells = table1.add_row().cells
        row_cells[0].text = libelle
        row_cells[1].text = valeur

    for i, width in enumerate(widths):
        for cell in table1.columns[i].cells:
            cell.width = Inches(width * 6)

    doc.add_paragraph()

    doc.add_paragraph("MONTANTS DE FRANCHISES (par sinistre exprimés en €)").runs[0].bold = True

    lignes2 = [
        ["Dommages subis par les ouvrages de bâtiment", context['franchise_ouvrage']],
        ["Catastrophes naturelles", "Montant déterminé par la loi ou par ses textes d'application"],
        ["Responsabilité civile (1)", ""],
        ["    - Assuré maître d'ouvrage", "xxxxxxxxxxxx"],
        ["    - Assuré maître d'ouvrage", "SANS"],
        ["Maintenance-visite", context['franchise_maintenance']],
    ]

    table2 = doc.add_table(rows=0, cols=2)
    table2.style = 'Table Grid'

    for libelle, valeur in lignes2:
        row_cells = table2.add_row().cells
        row_cells[0].text = libelle
        row_cells[1].text = valeur

    for i, width in enumerate(widths):
        for cell in table2.columns[i].cells:
            cell.width = Inches(width * 6)

    doc.add_paragraph()
    doc.add_paragraph("        (1) Ces franchises s’appliquent pour des dommages autres que corporels")
    doc.add_paragraph(f"Date de simulation de tarif : le {context['date_simulation']}")
    doc.add_paragraph("(*) Cette tarification est faite sous réserve d’acceptation du risque par la compagnie.")

    # Sauvegarde
    path = "attestation.docx"
    doc.save(path)
    return path
